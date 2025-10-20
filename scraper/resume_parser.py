"""
Resume parser for extracting text from PDF and DOCX files
"""
import re
from typing import Optional, Dict, Any
from pathlib import Path
import PyPDF2
import docx


class ResumeParser:
    """Parses resume files (PDF, DOCX) to extract text and basic information"""

    def parse_file(self, file_path: str, file_type: str) -> Optional[Dict[str, Any]]:
        """
        Parse a resume file and extract text

        Args:
            file_path: Path to the resume file
            file_type: File type ('pdf' or 'docx')

        Returns:
            Dictionary with extracted data, or None if parsing fails
        """
        try:
            # Extract raw text based on file type
            if file_type.lower() == 'pdf':
                raw_text = self._extract_text_from_pdf(file_path)
            elif file_type.lower() in ['docx', 'doc']:
                raw_text = self._extract_text_from_docx(file_path)
            else:
                print(f"Unsupported file type: {file_type}")
                return None

            if not raw_text or len(raw_text.strip()) < 50:
                print(f"Insufficient text extracted from {file_path}")
                return None

            # Clean raw text to remove null bytes and problematic characters
            raw_text = self._clean_text(raw_text)

            # Extract basic information
            extracted_data = {
                "raw_text": raw_text,
                "email": self._extract_email(raw_text),
                "phone": self._extract_phone(raw_text),
                "file_name": Path(file_path).name,
                "file_type": file_type,
            }

            return extracted_data

        except Exception as e:
            print(f"Error parsing file {file_path}: {e}")
            return None

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF file

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text as string
        """
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"Error extracting text from PDF {file_path}: {e}")

        return text.strip()

    def _extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text as string
        """
        text = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"

        except Exception as e:
            print(f"Error extracting text from DOCX {file_path}: {e}")

        return text.strip()

    def _extract_email(self, text: str) -> Optional[str]:
        """
        Extract email address from text

        Args:
            text: Text to search

        Returns:
            First email address found, or None
        """
        # Email regex pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        matches = re.findall(email_pattern, text)

        if matches:
            return matches[0]
        return None

    def _extract_phone(self, text: str) -> Optional[str]:
        """
        Extract phone number from text

        Args:
            text: Text to search

        Returns:
            First phone number found, or None
        """
        # Phone number patterns (US and international formats)
        phone_patterns = [
            r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',  # 123-456-7890 or 123.456.7890
            r'\(\d{3}\)\s*\d{3}[-.\s]?\d{4}',       # (123) 456-7890
            r'\+\d{1,3}\s*\d{1,4}\s*\d{1,4}\s*\d{1,4}',  # +1 123 456 7890
        ]

        for pattern in phone_patterns:
            matches = re.findall(pattern, text)
            if matches:
                return matches[0]

        return None

    def _clean_text(self, text: str) -> str:
        """
        Clean text by removing null bytes and other problematic characters

        Args:
            text: Raw text to clean

        Returns:
            Cleaned text safe for database storage
        """
        # Remove null bytes (\x00 or \u0000) - PostgreSQL can't handle these
        text = text.replace('\x00', '')
        text = text.replace('\u0000', '')

        # Remove other control characters except newlines and tabs
        cleaned = ""
        for char in text:
            # Keep printable characters, newlines, tabs, and spaces
            if char.isprintable() or char in ['\n', '\r', '\t']:
                cleaned += char

        return cleaned

    def extract_skills(self, text: str, skill_list: list = None) -> list:
        """
        Extract skills from text based on a predefined skill list

        Args:
            text: Text to search
            skill_list: List of skills to look for (case-insensitive)

        Returns:
            List of found skills
        """
        if skill_list is None:
            # Default common tech skills
            skill_list = [
                "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust",
                "react", "angular", "vue", "node.js", "django", "flask", "fastapi",
                "aws", "azure", "gcp", "docker", "kubernetes", "terraform",
                "postgresql", "mysql", "mongodb", "redis",
                "machine learning", "deep learning", "nlp", "computer vision",
                "git", "ci/cd", "agile", "scrum"
            ]

        text_lower = text.lower()
        found_skills = []

        for skill in skill_list:
            if skill.lower() in text_lower:
                found_skills.append(skill)

        return list(set(found_skills))  # Remove duplicates


# Global parser instance
resume_parser = ResumeParser()
