"""
User Resume API routes
Handles user resume uploads
"""
from fastapi import APIRouter, UploadFile, File, HTTPException, status, Depends
from pathlib import Path
import uuid

from api.auth import get_user_id
from services.storage_service import storage_service
from services.user_resume_service import user_resume_service
from models.user_resumes import (
    UploadResponse,
    ListResumesResponse,
    UserResumeItem,
    RenameResumeRequest,
    RenameResumeResponse,
    DeleteResumeResponse
)
from config import supabase

router = APIRouter()


@router.post("/upload", response_model=UploadResponse)
async def upload_user_resume(
    file: UploadFile = File(...),
    user_id: str = Depends(get_user_id)
):
    """
    Upload and save user's resume to their profile

    Requires authentication via Clerk JWT token

    Args:
        file: Resume file (PDF, DOCX, DOC, or TXT)
        user_id: User ID from Clerk JWT (injected by dependency)

    Returns:
        UploadResponse with success status and file URL
    """
    try:
        # Validate file type
        allowed_extensions = {'.pdf', '.docx', '.doc', '.txt'}
        file_extension = Path(file.filename).suffix.lower()

        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File type {file_extension} not supported. Please upload PDF, DOCX, DOC, or TXT"
            )

        # Read file content
        file_content = await file.read()

        # Generate unique resume ID
        resume_id = str(uuid.uuid4())

        # Generate storage path: {user_id}/{resume_id}/original{ext}
        storage_path = f"{user_id}/{resume_id}/original{file_extension}"

        # Determine content type
        content_types = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".doc": "application/msword",
            ".txt": "text/plain",
        }
        content_type = content_types.get(file_extension, "application/octet-stream")

        # Upload to Supabase Storage using storage service
        file_url = storage_service.upload_file(
            bucket_name="user-resumes",
            storage_path=storage_path,
            file_content=file_content,
            content_type=content_type
        )

        # Insert into user_resumes table
        file_type = file_extension.lstrip('.')  # Remove leading dot

        supabase.table("user_resumes").insert({
            "id": resume_id,
            "user_id": user_id,
            "filename": file.filename,
            "file_url": file_url,
            "storage_path": storage_path,
            "file_type": file_type,
            "resume_source": "upload"  # Mark as uploaded resume
        }).execute()

        return UploadResponse(
            success=True,
            message="Resume uploaded successfully",
            resume_id=resume_id,
            file_url=file_url
        )

    except HTTPException:
        raise
    except Exception as e:
        print(f"Error uploading resume: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to upload resume: {str(e)}"
        )


@router.get("/list", response_model=ListResumesResponse)
async def list_user_resumes(
    user_id: str = Depends(get_user_id)
):
    """
    Get all resumes for the authenticated user

    Returns all original resumes uploaded by the user.
    Does NOT include reviewed or anonymized versions (use separate endpoints for those).

    Args:
        user_id: User ID from Clerk JWT (injected by dependency)

    Returns:
        ListResumesResponse with list of user's resumes
    """
    try:
        # Get all resumes for this user
        result = supabase.table("user_resumes")\
            .select("id, filename, file_url, file_type, resume_source, builder_content, created_at, updated_at")\
            .eq("user_id", user_id)\
            .order("created_at", desc=True)\
            .execute()

        # Convert to Pydantic models
        resumes = [UserResumeItem(**resume) for resume in result.data]

        return ListResumesResponse(
            success=True,
            resumes=resumes
        )

    except Exception as e:
        print(f"Error listing resumes: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to list resumes: {str(e)}"
        )


@router.delete("/{resume_id}", response_model=DeleteResumeResponse)
async def delete_user_resume(
    resume_id: str,
    user_id: str = Depends(get_user_id)
):
    """
    Delete a user resume (works for both uploaded and builder resumes)

    Deletes the resume file from storage and removes the database record.

    Args:
        resume_id: UUID of the resume to delete
        user_id: User ID from Clerk JWT (injected by dependency)

    Returns:
        DeleteResumeResponse with success status
    """
    try:
        result = user_resume_service.delete_resume(resume_id, user_id)

        if not result.get("success"):
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=result.get("error", "Failed to delete resume")
            )

        return DeleteResumeResponse(
            success=True,
            message=result.get("message", "Resume deleted successfully")
        )

    except HTTPException:
        raise
    except Exception as e:
        print(f"Error deleting resume: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to delete resume: {str(e)}"
        )


@router.patch("/{resume_id}", response_model=RenameResumeResponse)
async def rename_user_resume(
    resume_id: str,
    request: RenameResumeRequest,
    user_id: str = Depends(get_user_id)
):
    """
    Rename a user resume (update filename)

    Updates the filename field in the database.

    Args:
        resume_id: UUID of the resume to rename
        request: Request body with new filename
        user_id: User ID from Clerk JWT (injected by dependency)

    Returns:
        RenameResumeResponse with success status and new filename
    """
    try:
        result = user_resume_service.rename_resume(
            resume_id, user_id, request.filename
        )

        if not result.get("success"):
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=result.get("error", "Failed to rename resume")
            )

        return RenameResumeResponse(
            success=True,
            message=result.get("message", "Resume renamed successfully"),
            resume_id=result.get("resume_id"),
            filename=result.get("filename")
        )

    except HTTPException:
        raise
    except Exception as e:
        print(f"Error renaming resume: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to rename resume: {str(e)}"
        )


def _extract_text_from_file(file_path: str, file_type: str) -> str:
    """
    Extract text from a file

    Args:
        file_path: Path to the file
        file_type: Type of file (pdf, docx, doc, txt)

    Returns:
        Extracted text
    """
    import PyPDF2
    import docx

    try:
        if file_type == 'pdf':
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip()

        elif file_type in ['docx', 'doc']:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            return text.strip()

        elif file_type == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        else:
            return ""

    except Exception as e:
        print(f"Error extracting text: {e}")
        return ""
